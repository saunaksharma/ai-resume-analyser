import pdfplumber


def extract_text(file_path: str) -> str:
    ext = file_path.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        return _extract_pdf(file_path)
    elif ext == "docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: .{ext}  (supported: pdf, docx)")


def _extract_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files. "
            "Install it with:  pip install python-docx"
        )

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)
